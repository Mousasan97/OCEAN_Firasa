"""
CV Parser Service - Extract career information from CV documents

Uses:
1. PyMuPDF (fitz) for PDF text extraction
2. python-docx for DOCX text extraction
3. LLM (via llm_provider) for structured data extraction
"""
import tempfile
import os
from typing import Optional
from pathlib import Path

from src.api.schemas.career import CareerProfile
from src.services.llm_provider import get_llm_provider
from src.utils.logger import get_logger

logger = get_logger(__name__)

CV_EXTRACTION_SYSTEM_PROMPT = """You are a CV/resume parser. Extract structured career information from the provided CV text.

Instructions:
- Extract only information that is explicitly stated in the CV
- For years_experience, calculate from work history dates (current year is 2025)
- For key_skills, list up to 10 most relevant technical and soft skills
- For industries, identify sectors the person has worked in
- Be concise and accurate
- If information is not found, use null

Return a JSON object with these fields:
- current_role: string (most recent job title) or null
- target_role: string (desired role if mentioned) or null
- location: string (city/country) or null
- years_experience: integer or null
- key_skills: array of strings (max 10)
- industries: array of strings
- education_level: string (e.g., "Bachelor's", "Master's", "PhD") or null
- certifications: array of strings
- summary: string (2-3 sentence professional summary) or null"""


class CVParserService:
    """Service for parsing CV documents and extracting career profiles."""

    def __init__(self):
        self.llm_provider = get_llm_provider()

    async def parse_cv(self, file_content: bytes, filename: str) -> CareerProfile:
        """
        Parse CV and extract career profile.

        Args:
            file_content: Raw bytes of the CV file
            filename: Original filename (used to determine file type)

        Returns:
            CareerProfile with extracted information
        """
        # Extract text based on file type
        ext = Path(filename).suffix.lower()

        if ext == '.pdf':
            text = self._extract_pdf_text(file_content)
        elif ext in ['.docx', '.doc']:
            text = self._extract_docx_text(file_content)
        else:
            raise ValueError(f"Unsupported file type: {ext}. Supported: .pdf, .docx, .doc")

        if not text or len(text.strip()) < 50:
            raise ValueError("Could not extract sufficient text from CV. Please ensure the file is not corrupted.")

        # Truncate if too long (LLM context limits)
        if len(text) > 15000:
            text = text[:15000] + "\n...[truncated]"

        logger.info(f"Extracted {len(text)} characters from CV: {filename}")

        # Use LLM to extract structured data
        profile = await self._extract_with_llm(text)
        return profile

    def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF using PyMuPDF."""
        try:
            import fitz  # PyMuPDF

            doc = fitz.open(stream=content, filetype="pdf")
            text_parts = []
            for page in doc:
                text_parts.append(page.get_text())
            doc.close()
            return "\n".join(text_parts)
        except ImportError:
            logger.error("PyMuPDF not installed. Run: pip install PyMuPDF")
            raise ValueError("PDF parsing not available. PyMuPDF not installed.")
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise ValueError(f"Failed to read PDF: {e}")

    def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX using python-docx."""
        try:
            from docx import Document

            # Save to temp file (python-docx needs file path)
            fd, tmp_path = tempfile.mkstemp(suffix='.docx')
            os.close(fd)
            try:
                with open(tmp_path, 'wb') as f:
                    f.write(content)
                doc = Document(tmp_path)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                paragraphs.append(cell.text.strip())

                return "\n".join(paragraphs)
            finally:
                os.unlink(tmp_path)
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            raise ValueError("DOCX parsing not available. python-docx not installed.")
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise ValueError(f"Failed to read DOCX: {e}")

    async def _extract_with_llm(self, cv_text: str) -> CareerProfile:
        """Use LLM to extract structured career profile from CV text."""
        user_prompt = f"Parse this CV/resume and extract the career information:\n\n{cv_text}"

        try:
            profile = await self.llm_provider.generate_structured_with_fallback(
                output_type=CareerProfile,
                system_prompt=CV_EXTRACTION_SYSTEM_PROMPT,
                user_prompt=user_prompt
            )

            logger.info(f"Successfully extracted career profile: role={profile.current_role}, location={profile.location}")
            return profile

        except Exception as e:
            logger.error(f"LLM extraction failed: {e}")
            raise ValueError(f"Failed to extract career information: {e}")


# Singleton
_cv_parser_service: Optional[CVParserService] = None


def get_cv_parser_service() -> CVParserService:
    """Get or create singleton CV parser service."""
    global _cv_parser_service
    if _cv_parser_service is None:
        _cv_parser_service = CVParserService()
    return _cv_parser_service
